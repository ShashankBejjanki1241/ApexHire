"""
Resume Parser Module
Extracts text from PDF and DOCX resume files
"""

import os
import logging

# Import centralized logging
from utils import setup_logging
logger = logging.getLogger(__name__)


class ResumeParser:
    """Parser for extracting text from resume files (PDF and DOCX)"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.doc', '.txt']
    
    def parse_resume(self, file_path: str):
        """Parse a resume file and extract text with comprehensive error handling"""
        try:
            # Validate file exists and is accessible
            if not os.path.exists(file_path):
                logger.error(f"File not found: {file_path}")
                return None
            
            if not os.access(file_path, os.R_OK):
                logger.error(f"File not readable: {file_path}")
                return None
            
            # Validate file size (max 50MB to prevent memory issues)
            file_size = os.path.getsize(file_path)
            max_size = 50 * 1024 * 1024  # 50MB
            if file_size > max_size:
                logger.error(f"File too large: {file_path} ({file_size / 1024 / 1024:.1f}MB > {max_size / 1024 / 1024}MB)")
                return None
            
            if file_size == 0:
                logger.error(f"File is empty: {file_path}")
                return None
            
            filename = os.path.basename(file_path)
            file_extension = os.path.splitext(filename)[1].lower()
            
            # Validate file extension
            if file_extension not in self.supported_formats:
                logger.error(f"Unsupported file format: {file_extension}. Supported: {', '.join(self.supported_formats)}")
                return None
            
            # Parse based on file type
            if file_extension == '.pdf':
                return self._parse_pdf(file_path)
            elif file_extension in ['.docx', '.doc']:
                return self._parse_docx(file_path)
            elif file_extension == '.txt':
                return self._parse_txt(file_path)
            else:
                logger.error(f"Unexpected file format: {file_extension}")
                return None
                
        except PermissionError as e:
            logger.error(f"Permission denied accessing {file_path}: {str(e)}")
            return None
        except OSError as e:
            logger.error(f"OS error accessing {file_path}: {str(e)}")
            return None
        except Exception as e:
            logger.error(f"Unexpected error parsing {file_path}: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            return None
    
    def _parse_pdf(self, file_path: str) -> str:
        """Parse PDF file and extract text with enhanced error handling"""
        try:
            import pdfplumber
            
            # Check if file is actually a valid PDF
            if not self._is_valid_pdf(file_path):
                logger.error(f"File is not a valid PDF: {file_path}")
                return None
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                # Validate PDF structure
                if not pdf.pages:
                    logger.error(f"PDF has no pages: {file_path}")
                    return None
                
                for i, page in enumerate(pdf.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
                        else:
                            logger.warning(f"Page {i+1} has no extractable text in {file_path}")
                    except Exception as e:
                        logger.warning(f"Error extracting text from page {i+1} in {file_path}: {str(e)}")
                        continue
            
            if not text.strip():
                logger.error(f"No text could be extracted from PDF: {file_path}")
                return None
            
            logger.info(f"Successfully parsed PDF: {file_path} ({len(text)} characters)")
            return text.strip()
            
        except ImportError:
            logger.error("pdfplumber not installed. Please install it: pip install pdfplumber")
            return None
        except Exception as e:
            logger.error(f"Error parsing PDF {file_path}: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            return None
    
    def _is_valid_pdf(self, file_path: str) -> bool:
        """Check if file is a valid PDF by reading header"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(4)
                return header == b'%PDF'
        except Exception:
            return False
    
    def _parse_docx(self, file_path: str) -> str:
        """Parse DOCX file and extract text with enhanced error handling"""
        try:
            from docx import Document
            
            # Check if file is actually a valid DOCX
            if not self._is_valid_docx(file_path):
                logger.error(f"File is not a valid DOCX: {file_path}")
                return None
            
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
            
            if not text.strip():
                logger.error(f"No text could be extracted from DOCX: {file_path}")
                return None
            
            logger.info(f"Successfully parsed DOCX: {file_path} ({len(text)} characters)")
            return text.strip()
            
        except ImportError:
            logger.error("python-docx not installed. Please install it: pip install python-docx")
            return None
        except Exception as e:
            logger.error(f"Error parsing DOCX {file_path}: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            return None
    
    def _is_valid_docx(self, file_path: str) -> bool:
        """Check if file is a valid DOCX by reading header"""
        try:
            import zipfile
            with zipfile.ZipFile(file_path, 'r') as zf:
                # Check for required DOCX files
                required_files = ['word/document.xml', 'word/_rels/document.xml.rels']
                return all(f in zf.namelist() for f in required_files)
        except Exception:
            return False
    
    def _parse_txt(self, file_path: str) -> str:
        """Parse TXT file and extract text with enhanced error handling"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            text = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        text = file.read()
                    logger.debug(f"Successfully read {file_path} with {encoding} encoding")
                    break
                except UnicodeDecodeError:
                    logger.debug(f"Failed to read {file_path} with {encoding} encoding")
                    continue
                except Exception as e:
                    logger.debug(f"Error reading {file_path} with {encoding} encoding: {str(e)}")
                    continue
            
            if text is None:
                logger.error(f"Could not read {file_path} with any supported encoding")
                return None
            
            if not text.strip():
                logger.error(f"TXT file is empty or contains only whitespace: {file_path}")
                return None
            
            logger.info(f"Successfully parsed TXT: {file_path} ({len(text)} characters)")
            return text.strip()
            
        except Exception as e:
            logger.error(f"Error parsing TXT {file_path}: {str(e)}")
            logger.error(f"Error type: {type(e).__name__}")
            return None
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file (alias for _parse_txt)"""
        return self._parse_txt(file_path)
    
    def parse_multiple_resumes(self, directory_path: str):
        """Parse multiple resume files from a directory"""
        parsed_resumes = []
        
        if not os.path.exists(directory_path):
            logger.error(f"Directory not found: {directory_path}")
            return parsed_resumes
        
        for filename in os.listdir(directory_path):
            file_path = os.path.join(directory_path, filename)
            
            if os.path.isfile(file_path):
                file_extension = os.path.splitext(filename)[1].lower()
                
                if file_extension in self.supported_formats:
                    parsed_resume = self.parse_resume(file_path)
                    if parsed_resume:
                        parsed_resumes.append(parsed_resume)
        
        logger.info(f"Successfully parsed {len(parsed_resumes)} resume files")
        return parsed_resumes


if __name__ == "__main__":
    parser = ResumeParser()
    print("Resume Parser Module")
